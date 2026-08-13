"""Convierte un .md a .docx con formato de Word, para pegar en el informe.

Traduce encabezados, negritas, código en línea, listas y tablas a estilos
nativos de Word, de modo que al copiar no arrastre la sintaxis de markdown.

Uso:
    python scripts/md_a_word.py docs/texto_informe_recuperacion.md
    python scripts/md_a_word.py <entrada.md> --out <salida.docx>
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# **negrita**, *cursiva* y `código`, en un solo barrido
_TROZOS = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def escribir_con_formato(parrafo, texto: str) -> None:
    """Añade `texto` al párrafo respetando negritas, cursivas y código."""
    for trozo in _TROZOS.split(texto):
        if not trozo:
            continue
        if trozo.startswith("**") and trozo.endswith("**"):
            parrafo.add_run(trozo[2:-2]).bold = True
        elif trozo.startswith("`") and trozo.endswith("`"):
            run = parrafo.add_run(trozo[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        elif trozo.startswith("*") and trozo.endswith("*"):
            parrafo.add_run(trozo[1:-1]).italic = True
        else:
            parrafo.add_run(trozo)


def es_separador_de_tabla(linea: str) -> bool:
    """La línea |---|---| que separa cabecera de cuerpo."""
    return bool(re.fullmatch(r"\|[\s:\-|]+\|", linea.strip()))


def celdas(linea: str) -> list[str]:
    return [c.strip() for c in linea.strip().strip("|").split("|")]


def agregar_tabla(doc: Document, filas: list[str]) -> None:
    cabecera = celdas(filas[0])
    cuerpo = [celdas(f) for f in filas[1:] if not es_separador_de_tabla(f)]
    tabla = doc.add_table(rows=1, cols=len(cabecera))
    tabla.style = "Light Grid Accent 1"
    for i, texto in enumerate(cabecera):
        celda = tabla.rows[0].cells[i]
        celda.text = ""
        escribir_con_formato(celda.paragraphs[0], texto)
        for run in celda.paragraphs[0].runs:
            run.bold = True
    for fila in cuerpo:
        nueva = tabla.add_row().cells
        for i, texto in enumerate(fila[:len(cabecera)]):
            nueva[i].text = ""
            escribir_con_formato(nueva[i].paragraphs[0], texto)
    doc.add_paragraph()


def convertir(md: Path, salida: Path) -> None:
    doc = Document()
    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)
    estilo.paragraph_format.space_after = Pt(8)

    lineas = md.read_text(encoding="utf-8").split("\n")
    i = 0
    buffer_parrafo: list[str] = []
    n_tablas = n_titulos = 0

    def volcar_parrafo():
        if buffer_parrafo:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            escribir_con_formato(p, " ".join(buffer_parrafo))
            buffer_parrafo.clear()

    while i < len(lineas):
        linea = lineas[i]
        desnuda = linea.strip()

        # Tabla: bloque de líneas consecutivas que empiezan por |
        if desnuda.startswith("|"):
            volcar_parrafo()
            filas = []
            while i < len(lineas) and lineas[i].strip().startswith("|"):
                filas.append(lineas[i])
                i += 1
            agregar_tabla(doc, filas)
            n_tablas += 1
            continue

        if not desnuda:
            volcar_parrafo()
        elif desnuda.startswith("#"):
            volcar_parrafo()
            nivel = len(desnuda) - len(desnuda.lstrip("#"))
            texto = desnuda.lstrip("#").strip()
            doc.add_heading(texto, level=min(nivel, 4))
            n_titulos += 1
        elif desnuda.startswith("---"):
            volcar_parrafo()          # separador: se ignora en Word
        elif desnuda.startswith(">"):
            volcar_parrafo()          # nota interna: no va al documento final
        elif re.match(r"^[-*]\s+", desnuda):
            volcar_parrafo()
            p = doc.add_paragraph(style="List Bullet")
            escribir_con_formato(p, re.sub(r"^[-*]\s+", "", desnuda))
        elif re.match(r"^\d+\.\s+", desnuda):
            volcar_parrafo()
            p = doc.add_paragraph(style="List Number")
            escribir_con_formato(p, re.sub(r"^\d+\.\s+", "", desnuda))
        else:
            buffer_parrafo.append(desnuda)
        i += 1

    volcar_parrafo()
    doc.save(salida)
    print(f"Escrito {salida}")
    print(f"  {n_titulos} encabezados, {n_tablas} tablas")


def main() -> None:
    p = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    p.add_argument("md", help="Archivo markdown de entrada.")
    p.add_argument("--out", default=None, help="Salida .docx (por defecto, el mismo nombre).")
    args = p.parse_args()
    entrada = Path(args.md)
    salida = Path(args.out) if args.out else entrada.with_suffix(".docx")
    convertir(entrada, salida)


if __name__ == "__main__":
    main()
